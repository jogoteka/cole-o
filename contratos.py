"""
Módulo de contratos de locação — integração com ClicksZap.

Fluxo:
  1. gerar_pdf_contrato(locacao_id) → bytes do PDF
     - Usa o modelo de cláusulas salvo no banco (tabela contrato_modelo)
     - Substitui todos os {{CAMPOS}} pelos dados reais da locação
  2. enviar_contrato(locacao_id)    → gera PDF + faz upload no ClicksZap
                                      + cria signature request (dispara WhatsApp)
                                      → salva token/status/request_id na locação
  3. status_contrato(locacao_id)    → consulta ClicksZap e retorna status + link download

Campos disponíveis para o template (use {{CAMPO}} no texto):
  {{NOME_LOJA}}        {{CNPJ_LOJA}}         {{ENDERECO_LOJA}}
  {{NUM_CONTRATO}}     {{DATA_GERACAO}}
  {{NOME_CLIENTE}}     {{CPF_CLIENTE}}        {{TELEFONE_CLIENTE}}
  {{ENDERECO_CLIENTE}}
  {{JOGO}}
  {{DATA_SAIDA}}       {{DATA_PREVISTA}}      {{OPCAO_DIAS}}
  {{VALOR_LOCACAO}}    {{FORMA_PAGAMENTO}}    {{MULTA_DIA}}
"""
import os
import io
import re
import logging
from datetime import datetime
from typing import Optional
from database import get_connection

logger = logging.getLogger(__name__)

# ── Configuração ──────────────────────────────────────────────────────────────

def _normalizar_url(url: str) -> str:
    """Remove www. para evitar 307 redirect que descarta o Authorization header."""
    url = url.rstrip("/")
    if url.startswith("https://www."):
        url = "https://" + url[len("https://www."):]
    elif url.startswith("http://www."):
        url = "http://" + url[len("http://www."):]
    return url

CLICKSZAP_URL = _normalizar_url(
    os.environ.get("CLICKSZAP_URL", "https://clickszap.com.br")
)

# Lido dinamicamente para pegar a env var mesmo após hot-reload
def _get_token() -> str:
    return os.environ.get("CLICKSZAP_TOKEN", "")

# Mantém alias para checagem rápida (não usar para o valor real)
CLICKSZAP_TOKEN = ""  # sempre via _get_token()

NOME_LOJA     = os.environ.get("NOME_LOJA",     "Jogoteka")
ENDERECO_LOJA = os.environ.get("ENDERECO_LOJA", "")
CNPJ_LOJA     = os.environ.get("CNPJ_LOJA",     "")

# ── Cláusulas padrão (usadas quando não há modelo salvo no banco) ─────────────

CLAUSULAS_PADRAO = """# Contrato de Locação de Jogo de Tabuleiro
## {{NOME_LOJA}} — Contrato #{{NUM_CONTRATO}}

1. LOCADORA
Nome: {{NOME_LOJA}}
CNPJ: {{CNPJ_LOJA}}
Endereço: {{ENDERECO_LOJA}}

2. LOCATÁRIO (CLIENTE)
Nome: {{NOME_CLIENTE}}
CPF: {{CPF_CLIENTE}}
Telefone / WhatsApp: {{TELEFONE_CLIENTE}}
Endereço: {{ENDERECO_CLIENTE}}

3. OBJETO DO CONTRATO
Jogo locado: {{JOGO}}

4. PRAZO E VALORES
Data de saída: {{DATA_SAIDA}}
Data prevista de devolução: {{DATA_PREVISTA}}
Período: {{OPCAO_DIAS}} dia(s)
Valor da locação: {{VALOR_LOCACAO}}
Forma de pagamento: {{FORMA_PAGAMENTO}}
Multa por dia de atraso: {{MULTA_DIA}}

5. CLÁUSULAS E CONDIÇÕES

5.1  O LOCATÁRIO recebe o jogo em perfeito estado e se compromete a devolvê-lo nas mesmas condições, com todas as peças, cartas e acessórios originais.

5.2  Em caso de devolução após a data prevista, será cobrada multa diária conforme o valor indicado na Cláusula 4, calculada sobre os dias de atraso.

5.3  Danos, perdas ou extravio de componentes do jogo são de responsabilidade exclusiva do LOCATÁRIO, que deverá arcar com o custo de reposição ou reparo.

5.4  É vedada a sublocação, empréstimo ou cessão do jogo a terceiros sem autorização expressa e por escrito da LOCADORA.

5.5  O presente contrato é regido pelas normas do Código Civil Brasileiro (Lei 10.406/2002) e legislação correlata.

6. ASSINATURAS

Ao assinar este documento eletronicamente, o LOCATÁRIO declara ter lido, compreendido e concordado com todas as cláusulas acima.

______________________________          ______________________________
{{NOME_LOJA}} — Locadora                {{NOME_CLIENTE}} — Locatário

Data de geração: {{DATA_GERACAO}}
Documento gerado por ClicksZap | www.clickszap.com.br"""


# ── Helpers de BD ─────────────────────────────────────────────────────────────

def _buscar_locacao(locacao_id: int) -> Optional[dict]:
    with get_connection() as conn:
        row = conn.execute("""
            SELECT l.*, j.nome AS jogo_nome, j.multa_dia,
                   c.nome AS cliente_nome, c.cpf AS cliente_cpf,
                   c.telefone AS cliente_tel, c.logradouro, c.numero,
                   c.bairro, c.cidade, c.estado
            FROM locacoes l
            JOIN jogos j ON j.id = l.jogo_id
            LEFT JOIN clientes c ON c.id = l.cliente_id
            WHERE l.id = ?
        """, (locacao_id,)).fetchone()
        return dict(row) if row else None


def _buscar_grupo_locacoes(locacao_id: int) -> list:
    """
    Retorna a lista de locações do mesmo cliente feitas na mesma data.
    Usado para gerar um contrato único cobrindo múltiplos jogos.
    A locação principal (locacao_id) é sempre a primeira da lista.
    """
    with get_connection() as conn:
        # Pega a locação principal para descobrir cliente_id e data
        principal = conn.execute(
            "SELECT cliente_id, data_saida FROM locacoes WHERE id = ?", (locacao_id,)
        ).fetchone()
        if not principal:
            return []

        cliente_id = principal["cliente_id"]
        data_saida = str(principal["data_saida"])[:10]  # apenas YYYY-MM-DD

        if not cliente_id:
            # Sem cliente identificado → retorna só a locação principal
            loc = _buscar_locacao(locacao_id)
            return [loc] if loc else []

        # Busca todas do mesmo cliente no mesmo dia
        rows = conn.execute("""
            SELECT l.*, j.nome AS jogo_nome, j.multa_dia,
                   c.nome AS cliente_nome, c.cpf AS cliente_cpf,
                   c.telefone AS cliente_tel, c.logradouro, c.numero,
                   c.bairro, c.cidade, c.estado
            FROM locacoes l
            JOIN jogos j ON j.id = l.jogo_id
            LEFT JOIN clientes c ON c.id = l.cliente_id
            WHERE l.cliente_id = ?
              AND substr(l.data_saida, 1, 10) = ?
            ORDER BY l.id
        """, (cliente_id, data_saida)).fetchall()

        locs = [dict(r) for r in rows]
        # Garante que a locação principal é a primeira
        locs.sort(key=lambda r: (0 if r["id"] == locacao_id else 1, r["id"]))
        return locs


def _salvar_dados_contrato(locacao_ids: list, request_id: str, token: str, status: str):
    """Salva token/status em todas as locações do grupo."""
    with get_connection() as conn:
        for lid in locacao_ids:
            conn.execute("""
                UPDATE locacoes
                SET contrato_request_id = ?,
                    contrato_token      = ?,
                    contrato_status     = ?
                WHERE id = ?
            """, (request_id, token, status, lid))


def _atualizar_status_contrato(locacao_id: int, status: str):
    with get_connection() as conn:
        conn.execute(
            "UPDATE locacoes SET contrato_status = ? WHERE id = ?",
            (status, locacao_id)
        )


def carregar_modelo() -> str:
    """Retorna o texto do modelo de cláusulas salvo no banco, ou o padrão."""
    try:
        with get_connection() as conn:
            row = conn.execute(
                "SELECT clausulas FROM contrato_modelo ORDER BY id DESC LIMIT 1"
            ).fetchone()
            if row and row["clausulas"]:
                return row["clausulas"]
    except Exception:
        pass
    return CLAUSULAS_PADRAO


def salvar_modelo(clausulas: str):
    """Salva (upsert) o modelo de cláusulas no banco."""
    agora = datetime.now().isoformat(sep=" ", timespec="seconds")
    with get_connection() as conn:
        existing = conn.execute("SELECT id FROM contrato_modelo LIMIT 1").fetchone()
        if existing:
            conn.execute(
                "UPDATE contrato_modelo SET clausulas = ?, atualizado_em = ? WHERE id = ?",
                (clausulas, agora, existing["id"])
            )
        else:
            conn.execute(
                "INSERT INTO contrato_modelo (id, clausulas, atualizado_em) VALUES (1, ?, ?)",
                (clausulas, agora)
            )


def salvar_template_pdf(pdf_bytes: bytes):
    """Salva o PDF de contrato padrão do usuário no banco (base64)."""
    import base64
    b64 = base64.b64encode(pdf_bytes).decode("utf-8")
    agora = datetime.now().isoformat(sep=" ", timespec="seconds")
    clausulas_atual = carregar_modelo()
    with get_connection() as conn:
        _garantir_coluna_docx(conn)
        existing = conn.execute("SELECT id FROM contrato_modelo LIMIT 1").fetchone()
        if existing:
            # Subir um PDF desativa o DOCX (o PDF passa a ser o modelo ativo)
            conn.execute(
                "UPDATE contrato_modelo SET template_pdf_b64 = ?, template_docx_b64 = NULL, atualizado_em = ? WHERE id = ?",
                (b64, agora, existing["id"])
            )
        else:
            conn.execute(
                "INSERT INTO contrato_modelo (id, clausulas, template_pdf_b64, atualizado_em) VALUES (1, ?, ?, ?)",
                (clausulas_atual, b64, agora)
            )


def carregar_template_pdf() -> Optional[bytes]:
    """Retorna os bytes do PDF de contrato padrão, ou None se não houver."""
    import base64
    try:
        with get_connection() as conn:
            row = conn.execute(
                "SELECT template_pdf_b64 FROM contrato_modelo ORDER BY id DESC LIMIT 1"
            ).fetchone()
            if row and row["template_pdf_b64"]:
                return base64.b64decode(row["template_pdf_b64"])
    except Exception:
        pass
    return None


def remover_template_pdf():
    """Remove o PDF de contrato padrão (volta a usar apenas a folha de dados)."""
    agora = datetime.now().isoformat(sep=" ", timespec="seconds")
    with get_connection() as conn:
        conn.execute(
            "UPDATE contrato_modelo SET template_pdf_b64 = NULL, atualizado_em = ?",
            (agora,)
        )


def _garantir_coluna_docx(conn):
    """Garante que a coluna template_docx_b64 existe — cria se não existir."""
    from database import DATABASE_URL, _get_cols
    try:
        cols = _get_cols(conn, "contrato_modelo")
        if "template_docx_b64" not in cols:
            add = "ADD COLUMN IF NOT EXISTS" if DATABASE_URL else "ADD COLUMN"
            conn.execute(f"ALTER TABLE contrato_modelo {add} template_docx_b64 TEXT")
    except Exception as e:
        logger.warning("_garantir_coluna_docx: %s", e)


def salvar_template_docx(docx_bytes: bytes):
    """Salva o DOCX de contrato padrão no banco (base64)."""
    import base64
    b64 = base64.b64encode(docx_bytes).decode("utf-8")
    agora = datetime.now().isoformat(sep=" ", timespec="seconds")
    with get_connection() as conn:
        _garantir_coluna_docx(conn)
        existing = conn.execute("SELECT id FROM contrato_modelo LIMIT 1").fetchone()
        if existing:
            conn.execute(
                "UPDATE contrato_modelo SET template_docx_b64 = ?, template_pdf_b64 = NULL, atualizado_em = ? WHERE id = ?",
                (b64, agora, existing["id"])
            )
        else:
            conn.execute(
                "INSERT INTO contrato_modelo (id, clausulas, template_docx_b64, atualizado_em) VALUES (1, ?, ?, ?)",
                (CLAUSULAS_PADRAO, b64, agora)
            )


def carregar_template_docx() -> Optional[bytes]:
    """Retorna os bytes do DOCX de contrato padrão, ou None se não houver."""
    import base64
    try:
        with get_connection() as conn:
            row = conn.execute(
                "SELECT template_docx_b64 FROM contrato_modelo ORDER BY id DESC LIMIT 1"
            ).fetchone()
            if row and row["template_docx_b64"]:
                return base64.b64decode(row["template_docx_b64"])
    except Exception:
        pass
    return None


def remover_template_docx():
    agora = datetime.now().isoformat(sep=" ", timespec="seconds")
    with get_connection() as conn:
        conn.execute(
            "UPDATE contrato_modelo SET template_docx_b64 = NULL, atualizado_em = ?",
            (agora,)
        )


def info_template() -> dict:
    """Retorna qual tipo de template está ativo: 'docx', 'pdf' ou 'texto'."""
    try:
        with get_connection() as conn:
            row = conn.execute(
                "SELECT template_docx_b64, template_pdf_b64 FROM contrato_modelo ORDER BY id DESC LIMIT 1"
            ).fetchone()
            if row:
                if row["template_docx_b64"]:
                    return {"tipo": "docx"}
                if row["template_pdf_b64"]:
                    return {"tipo": "pdf"}
    except Exception:
        pass
    return {"tipo": "texto"}


# Nomes de campos que o sistema sabe preencher
CAMPOS_DISPONIVEIS = [
    "NOME_LOJA", "CNPJ_LOJA", "ENDERECO_LOJA", "NUM_CONTRATO", "DATA_GERACAO",
    "NOME_CLIENTE", "CPF_CLIENTE", "TELEFONE_CLIENTE", "ENDERECO_CLIENTE",
    "JOGO", "DATA_SAIDA", "DATA_PREVISTA", "OPCAO_DIAS", "VALOR_LOCACAO",
    "FORMA_PAGAMENTO", "MULTA_DIA", "TOTAL_JOGOS", "VALOR_TOTAL",
] + [f"{base}_{i}" for i in range(1, 6)
     for base in ("JOGO", "DATA_SAIDA", "DATA_PREVISTA", "OPCAO_DIAS",
                  "VALOR_LOCACAO", "FORMA_PAGAMENTO", "MULTA_DIA")]


def diagnostico_campos_pdf() -> dict:
    """
    Lista os campos de formulário do PDF ativo e indica quais o sistema
    consegue preencher. Útil para o usuário conferir os nomes dos campos.
    """
    pdf_bytes = carregar_template_pdf()
    if not pdf_bytes:
        return {"tem_pdf": False, "campos": [], "msg": "Nenhum PDF cadastrado como modelo ativo."}

    from pypdf import PdfReader
    try:
        reader = PdfReader(io.BytesIO(pdf_bytes))
        fields = reader.get_fields() or {}
    except Exception as e:
        return {"tem_pdf": True, "campos": [], "erro": f"Erro ao ler o PDF: {e}"}

    validos = {c.lower() for c in CAMPOS_DISPONIVEIS}
    campos = []
    for nome in fields.keys():
        chave = (nome or "").strip().lower().replace("{{", "").replace("}}", "")
        campos.append({"nome": nome, "reconhecido": chave in validos})

    n_ok = sum(1 for c in campos if c["reconhecido"])
    return {
        "tem_pdf": True,
        "total_campos": len(campos),
        "reconhecidos": n_ok,
        "campos": campos,
        "campos_disponiveis": CAMPOS_DISPONIVEIS,
    }


def _docx_para_pdf(docx_bytes: bytes, locs) -> bytes:
    """
    1. Abre o DOCX com python-docx
    2. Substitui {{CAMPOS}} (incluindo numerados) em todos os parágrafos e tabelas
    3. Extrai o conteúdo (parágrafos, tabelas e imagens) e renderiza com ReportLab
    """
    import docx as _docx
    from docx.oxml.ns import qn
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                    HRFlowable, Table as RLTable, TableStyle,
                                    Image as RLImage)
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY, TA_RIGHT

    campos = _build_campos(locs)
    loc = locs[0] if isinstance(locs, list) else locs  # para rodapé

    # ── Abre e substitui no DOCX ──────────────────────────────────────────────
    doc = _docx.Document(io.BytesIO(docx_bytes))

    def _subst(texto: str) -> str:
        for campo, valor in campos.items():
            texto = texto.replace("{{" + campo + "}}", valor)
        return texto

    def _subst_para(para):
        """Substitui {{CAMPOS}} num parágrafo preservando a formatação básica."""
        full = "".join(r.text for r in para.runs)
        new  = _subst(full)
        if new != full and para.runs:
            para.runs[0].text = new
            for r in para.runs[1:]:
                r.text = ""

    for para in doc.paragraphs:
        _subst_para(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _subst_para(para)
    for section in doc.sections:
        for para in section.header.paragraphs:
            _subst_para(para)
        for para in section.footer.paragraphs:
            _subst_para(para)

    # ── Estilos ReportLab ─────────────────────────────────────────────────────
    dark  = HexColor("#1a1a2e")
    muted = HexColor("#6c757d")
    txt   = HexColor("#212529")
    max_img_w = A4[0] - 5*cm  # largura máxima de imagem (respeitando margens)

    styles = getSampleStyleSheet()
    s_normal = ParagraphStyle("cn", parent=styles["Normal"],
        fontName="Helvetica", fontSize=10, leading=15,
        textColor=txt, spaceAfter=4, alignment=TA_JUSTIFY)
    s_h1 = ParagraphStyle("ch1", parent=styles["Normal"],
        fontName="Helvetica-Bold", fontSize=16, leading=20,
        textColor=dark, spaceBefore=6, spaceAfter=4, alignment=TA_CENTER)
    s_h2 = ParagraphStyle("ch2", parent=styles["Normal"],
        fontName="Helvetica-Bold", fontSize=13, leading=17,
        textColor=dark, spaceBefore=10, spaceAfter=3)
    s_h3 = ParagraphStyle("ch3", parent=styles["Normal"],
        fontName="Helvetica-Bold", fontSize=11, leading=15,
        textColor=dark, spaceBefore=8, spaceAfter=2)
    s_bold = ParagraphStyle("cbold", parent=s_normal,
        fontName="Helvetica-Bold")
    s_center = ParagraphStyle("ccenter", parent=s_normal,
        alignment=TA_CENTER)

    # Mapeamento de estilo Word → ReportLab
    HEADING_STYLES = {
        "Heading 1": s_h1, "Título 1": s_h1,
        "Heading 2": s_h2, "Título 2": s_h2,
        "Heading 3": s_h3, "Título 3": s_h3,
    }

    def _extract_images(para):
        """Extrai imagens inline do parágrafo como lista de RLImage."""
        imgs = []
        for drawing in para._element.findall('.//' + qn('w:drawing')):
            for blip in drawing.findall('.//' + qn('a:blip')):
                rId = blip.get(
                    '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed'
                )
                if not rId:
                    continue
                try:
                    img_bytes = doc.part.related_parts[rId].blob
                except (KeyError, AttributeError):
                    continue
                # Tenta ler dimensões originais do Word (em EMUs: 1 cm = 360000 EMU)
                w_rl = h_rl = None
                extent = drawing.find('.//' + qn('wp:extent'))
                if extent is not None:
                    cx = int(extent.get('cx', 0))
                    cy = int(extent.get('cy', 0))
                    if cx > 0 and cy > 0:
                        w_rl = cx / 360000 * cm
                        h_rl = cy / 360000 * cm
                        if w_rl > max_img_w:
                            scale = max_img_w / w_rl
                            w_rl *= scale
                            h_rl *= scale
                try:
                    if w_rl and h_rl:
                        imgs.append(RLImage(io.BytesIO(img_bytes), width=w_rl, height=h_rl))
                    else:
                        imgs.append(RLImage(io.BytesIO(img_bytes), width=min(6*cm, max_img_w), kind="proportional"))
                except Exception:
                    pass
        return imgs

    def _para_to_rl_list(para):
        """Converte um parágrafo do DOCX para uma lista de elementos ReportLab (texto + imagens)."""
        elements = []

        # Imagens inline primeiro
        elements.extend(_extract_images(para))

        texto = para.text.strip()
        if texto:
            style_name = para.style.name if para.style else ""
            if style_name in HEADING_STYLES:
                elements.append(Paragraph(texto, HEADING_STYLES[style_name]))
            else:
                bold_runs = sum(1 for r in para.runs if r.bold and r.text.strip())
                all_runs  = sum(1 for r in para.runs if r.text.strip())
                if all_runs and bold_runs / all_runs > 0.5:
                    elements.append(Paragraph(texto, s_bold))
                else:
                    elements.append(Paragraph(texto, s_normal))
        elif not elements:
            elements.append(Spacer(1, 6))

        return elements

    def _cell_content(cell):
        """Retorna conteúdo de uma célula de tabela: texto e/ou imagens."""
        from reportlab.platypus import KeepTogether
        s_cell = ParagraphStyle("ct", parent=s_normal, fontSize=9, leading=12)
        elementos = []
        for para in cell.paragraphs:
            elementos.extend(_extract_images(para))
            txt = para.text.strip()
            if txt:
                elementos.append(Paragraph(txt, s_cell))
        if not elementos:
            return Paragraph("", s_cell)
        if len(elementos) == 1:
            return elementos[0]
        return KeepTogether(elementos)

    def _table_to_rl(table):
        """Converte uma tabela do DOCX para uma RLTable do ReportLab (com imagens em células)."""
        data = []
        for row in table.rows:
            linha = [_cell_content(cell) for cell in row.cells]
            data.append(linha)
        if not data:
            return None
        col_w = [(A4[0] - 4*cm) / len(data[0])] * len(data[0])
        t = RLTable(data, colWidths=col_w, repeatRows=1)
        t.setStyle(TableStyle([
            ("GRID",       (0,0), (-1,-1), 0.5, HexColor("#dee2e6")),
            ("BACKGROUND", (0,0), (-1,0),  HexColor("#f8f9fa")),
            ("FONTNAME",   (0,0), (-1,0),  "Helvetica-Bold"),
            ("FONTSIZE",   (0,0), (-1,-1), 9),
            ("VALIGN",     (0,0), (-1,-1), "TOP"),
            ("TOPPADDING", (0,0), (-1,-1), 4),
            ("BOTTOMPADDING", (0,0), (-1,-1), 4),
        ]))
        return t

    # ── Monta o story ─────────────────────────────────────────────────────────
    buf = io.BytesIO()
    pdf_doc = SimpleDocTemplate(buf, pagesize=A4,
        leftMargin=2.5*cm, rightMargin=2.5*cm,
        topMargin=2.5*cm, bottomMargin=2.5*cm)

    story = []

    # Imagens do cabeçalho do Word (logo, etc.) aparecem no topo do PDF
    for section in doc.sections:
        for para in section.header.paragraphs:
            imgs = _extract_images(para)
            story.extend(imgs)
        if story:
            story.append(Spacer(1, 8))
            break  # usa só o primeiro cabeçalho

    # Itera pelo corpo do documento na ordem (parágrafos e tabelas intercalados)
    body = doc.element.body
    para_idx = 0
    tbl_idx  = 0
    all_paras  = doc.paragraphs
    all_tables = doc.tables

    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "p" and para_idx < len(all_paras):
            for elem in _para_to_rl_list(all_paras[para_idx]):
                story.append(elem)
            para_idx += 1
        elif tag == "tbl" and tbl_idx < len(all_tables):
            elem = _table_to_rl(all_tables[tbl_idx])
            if elem:
                story.append(Spacer(1, 6))
                story.append(elem)
                story.append(Spacer(1, 6))
            tbl_idx += 1

    if not story:
        story.append(Paragraph("Contrato sem conteúdo", s_normal))

    def _rodape(canvas, doc):
        canvas.saveState()
        canvas.setFont("Helvetica-Oblique", 6.5)
        canvas.setFillColor(muted)
        canvas.drawString(2.5*cm, 1.2*cm,
            f"Contrato #{loc['id']:05d}  |  {NOME_LOJA}  |  "
            "Gerado por ClicksZap  |  www.clickszap.com.br")
        canvas.drawRightString(A4[0] - 2.5*cm, 1.2*cm, f"Página {doc.page}")
        canvas.restoreState()

    pdf_doc.build(story, onFirstPage=_rodape, onLaterPages=_rodape)
    return buf.getvalue()


def _build_campos(locs) -> dict:
    """
    Monta o dicionário de campos para substituição.
    Aceita uma única locação (dict) ou uma lista de locações.
    Gera campos numerados: JOGO_1..5, VALOR_LOCACAO_1..5, etc.
    Campos sem jogo correspondente ficam vazios ("").
    """
    if isinstance(locs, dict):
        locs = [locs]

    def fmt_data(d):
        try:
            return datetime.strptime(str(d)[:10], "%Y-%m-%d").strftime("%d/%m/%Y")
        except Exception:
            return str(d) if d else "—"

    def fmt_moeda(v):
        try:
            return "R$ {:,.2f}".format(float(v)).replace(",", "X").replace(".", ",").replace("X", ".")
        except Exception:
            return "—"

    pri = locs[0]  # locação principal (dados do cliente, loja, etc.)

    endereco_cli = ", ".join(filter(None, [
        pri.get("logradouro"), pri.get("numero"),
        pri.get("bairro"), pri.get("cidade"), pri.get("estado")
    ])) or "Não informado"

    # Campos base (cliente + loja)
    campos = {
        "NOME_LOJA":        NOME_LOJA,
        "CNPJ_LOJA":        CNPJ_LOJA or "Não informado",
        "ENDERECO_LOJA":    ENDERECO_LOJA or "Não informado",
        "NUM_CONTRATO":     f"{pri['id']:05d}",
        "DATA_GERACAO":     datetime.now().strftime("%d/%m/%Y %H:%M"),
        "NOME_CLIENTE":     pri.get("cliente_nome") or "Não identificado",
        "CPF_CLIENTE":      pri.get("cliente_cpf") or "Não informado",
        "TELEFONE_CLIENTE": pri.get("cliente_tel") or "Não informado",
        "ENDERECO_CLIENTE": endereco_cli,
        # Campos do primeiro jogo (compatibilidade com contratos de 1 jogo)
        "JOGO":             pri.get("jogo_nome") or "—",
        "DATA_SAIDA":       fmt_data(pri.get("data_saida")),
        "DATA_PREVISTA":    fmt_data(pri.get("data_prevista")),
        "OPCAO_DIAS":       str(pri.get("opcao_dias") or "—"),
        "VALOR_LOCACAO":    fmt_moeda(pri.get("valor_locacao") or 0),
        "FORMA_PAGAMENTO":  pri.get("forma_pagamento") or "Não informado",
        "MULTA_DIA":        (fmt_moeda(pri.get("multa_dia")) + "/dia") if pri.get("multa_dia") else "Sem multa cadastrada",
        "TOTAL_JOGOS":      str(len(locs)),
    }

    # Campos numerados: JOGO_1 a JOGO_5
    valor_total = 0.0
    for i in range(1, 6):
        if i <= len(locs):
            l = locs[i - 1]
            multa = l.get("multa_dia") or 0
            val = float(l.get("valor_locacao") or 0)
            valor_total += val
            campos[f"JOGO_{i}"]             = l.get("jogo_nome") or "—"
            campos[f"DATA_SAIDA_{i}"]       = fmt_data(l.get("data_saida"))
            campos[f"DATA_PREVISTA_{i}"]    = fmt_data(l.get("data_prevista"))
            campos[f"OPCAO_DIAS_{i}"]       = str(l.get("opcao_dias") or "—")
            campos[f"VALOR_LOCACAO_{i}"]    = fmt_moeda(val)
            campos[f"FORMA_PAGAMENTO_{i}"]  = l.get("forma_pagamento") or "Não informado"
            campos[f"MULTA_DIA_{i}"]        = (fmt_moeda(multa) + "/dia") if multa else "Sem multa"
        else:
            # Slot vazio — apaga o campo do template
            campos[f"JOGO_{i}"]             = ""
            campos[f"DATA_SAIDA_{i}"]       = ""
            campos[f"DATA_PREVISTA_{i}"]    = ""
            campos[f"OPCAO_DIAS_{i}"]       = ""
            campos[f"VALOR_LOCACAO_{i}"]    = ""
            campos[f"FORMA_PAGAMENTO_{i}"]  = ""
            campos[f"MULTA_DIA_{i}"]        = ""

    campos["VALOR_TOTAL"] = fmt_moeda(valor_total)
    return campos


def _substituir_campos(texto: str, locs) -> str:
    """Substitui todos os {{CAMPOS}} no texto pelo valor real da(s) locação(ões)."""
    campos = _build_campos(locs)
    resultado = texto
    for campo, valor in campos.items():
        resultado = resultado.replace("{{" + campo + "}}", valor)
    return resultado


# ── Geração do PDF ────────────────────────────────────────────────────────────

def _gerar_folha_dados(locs) -> bytes:
    """
    Gera uma página de dados da(s) locação(ões) (capa) usando ReportLab.
    Suporta múltiplos jogos do mesmo cliente no mesmo dia.
    """
    if isinstance(locs, dict):
        locs = [locs]
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable, Table, TableStyle, Image as RLImage
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_LEFT

    dark   = HexColor("#1a1a2e")
    accent = HexColor("#ED940E")
    muted  = HexColor("#6c757d")
    txt    = HexColor("#212529")
    green  = HexColor("#17C629")

    def fmt_data(d):
        try:
            return datetime.strptime(str(d)[:10], "%Y-%m-%d").strftime("%d/%m/%Y")
        except Exception:
            return str(d) if d else "—"

    def fmt_moeda(v):
        try:
            return "R$ {:,.2f}".format(float(v)).replace(",", "X").replace(".", ",").replace("X", ".")
        except Exception:
            return "—"

    loc = locs[0]  # locação principal (dados do cliente)

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
        leftMargin=2*cm, rightMargin=2*cm,
        topMargin=2*cm, bottomMargin=2*cm)

    styles = getSampleStyleSheet()
    h_loja = ParagraphStyle("h_loja", parent=styles["Normal"],
        fontName="Helvetica-Bold", fontSize=22, textColor=accent,
        spaceBefore=0, spaceAfter=2, alignment=TA_CENTER)
    h_sub = ParagraphStyle("h_sub", parent=styles["Normal"],
        fontName="Helvetica", fontSize=10, textColor=muted,
        spaceBefore=0, spaceAfter=0, alignment=TA_CENTER)
    h_num = ParagraphStyle("h_num", parent=styles["Normal"],
        fontName="Helvetica-Bold", fontSize=9, textColor=muted,
        spaceBefore=0, spaceAfter=0, alignment=TA_CENTER)
    label_s = ParagraphStyle("label", parent=styles["Normal"],
        fontName="Helvetica-Bold", fontSize=7.5, textColor=muted,
        spaceBefore=0, spaceAfter=1)
    value_s = ParagraphStyle("value", parent=styles["Normal"],
        fontName="Helvetica-Bold", fontSize=11, textColor=dark,
        spaceBefore=0, spaceAfter=0)
    value_sm = ParagraphStyle("value_sm", parent=styles["Normal"],
        fontName="Helvetica", fontSize=9.5, textColor=txt,
        spaceBefore=0, spaceAfter=0)

    story = []
    _logo = os.path.join(os.path.dirname(os.path.abspath(__file__)), "jogoteka_colorido.png")
    if os.path.exists(_logo):
        _img = RLImage(_logo, width=6*cm, height=3*cm, kind="proportional")
        _img.hAlign = "CENTER"
        story.append(_img)
        story.append(Spacer(1, 4))
    else:
        story.append(Paragraph(NOME_LOJA, h_loja))
    story.append(Paragraph(f"Contrato #{loc['id']:05d}  ·  Gerado em {datetime.now().strftime('%d/%m/%Y %H:%M')}", h_num))
    story.append(HRFlowable(width="100%", thickness=2, color=accent, spaceAfter=12, spaceBefore=8))

    def bloco(rotulo, valor, style=value_s):
        return [Paragraph(rotulo.upper(), label_s), Paragraph(str(valor), style), Spacer(1, 8)]

    endereco_cli = ", ".join(filter(None, [
        loc.get("logradouro"), loc.get("numero"),
        loc.get("bairro"), loc.get("cidade"), loc.get("estado")
    ])) or "Não informado"

    # Seção Locatário
    story.append(Paragraph("LOCATÁRIO", ParagraphStyle("sec", parent=styles["Normal"],
        fontName="Helvetica-Bold", fontSize=8, textColor=accent,
        spaceBefore=0, spaceAfter=4, borderPad=4)))
    data_table = [
        [Paragraph("NOME", label_s), Paragraph("CPF", label_s), Paragraph("TELEFONE / WHATSAPP", label_s)],
        [Paragraph(loc.get("cliente_nome") or "—", value_s),
         Paragraph(loc.get("cliente_cpf") or "—", value_sm),
         Paragraph(loc.get("cliente_tel") or "—", value_sm)],
    ]
    t = Table(data_table, colWidths=["45%", "25%", "30%"])
    t.setStyle(TableStyle([
        ("VALIGN", (0,0), (-1,-1), "TOP"),
        ("BOTTOMPADDING", (0,0), (-1,-1), 4),
    ]))
    story.append(t)
    story.append(Spacer(1, 4))
    story.extend(bloco("Endereço", endereco_cli, value_sm))
    story.append(HRFlowable(width="100%", thickness=0.4, color=HexColor("#dee2e6"), spaceAfter=8, spaceBefore=4))

    # Seção Jogos (suporta múltiplos)
    titulo_jogos = "JOGOS LOCADOS" if len(locs) > 1 else "JOGO LOCADO"
    story.append(Paragraph(titulo_jogos, ParagraphStyle("sec2", parent=styles["Normal"],
        fontName="Helvetica-Bold", fontSize=8, textColor=accent, spaceBefore=0, spaceAfter=4)))

    # Tabela de jogos: uma linha por jogo
    hdr_jogos = [Paragraph("JOGO", label_s), Paragraph("SAÍDA", label_s),
                 Paragraph("DEVOLUÇÃO", label_s), Paragraph("DIAS", label_s),
                 Paragraph("VALOR", label_s), Paragraph("PAGAMENTO", label_s)]
    rows_jogos = [hdr_jogos]
    valor_total = 0.0
    for l in locs:
        val = float(l.get("valor_locacao") or 0)
        valor_total += val
        rows_jogos.append([
            Paragraph(l.get("jogo_nome") or "—", value_sm),
            Paragraph(fmt_data(l.get("data_saida")), value_sm),
            Paragraph(fmt_data(l.get("data_prevista")), value_sm),
            Paragraph(f"{l.get('opcao_dias') or '—'}d", value_sm),
            Paragraph(fmt_moeda(val), value_sm),
            Paragraph(l.get("forma_pagamento") or "—", value_sm),
        ])
    if len(locs) > 1:
        rows_jogos.append([
            Paragraph("TOTAL", ParagraphStyle("tot", parent=label_s, fontName="Helvetica-Bold")),
            Paragraph("", value_sm), Paragraph("", value_sm), Paragraph("", value_sm),
            Paragraph(fmt_moeda(valor_total), value_s),
            Paragraph("", value_sm),
        ])

    tbl_jogos = Table(rows_jogos, colWidths=["28%","13%","13%","8%","16%","22%"])
    tbl_jogos.setStyle(TableStyle([
        ("GRID",       (0,0),(-1,-1), 0.3, HexColor("#dee2e6")),
        ("BACKGROUND", (0,0),(-1,0),  HexColor("#f8f9fa")),
        ("FONTNAME",   (0,0),(-1,0),  "Helvetica-Bold"),
        ("FONTSIZE",   (0,0),(-1,-1), 8),
        ("VALIGN",     (0,0),(-1,-1), "TOP"),
        ("TOPPADDING", (0,0),(-1,-1), 3),
        ("BOTTOMPADDING",(0,0),(-1,-1), 3),
    ]))
    story.append(tbl_jogos)
    story.append(HRFlowable(width="100%", thickness=0.4, color=HexColor("#dee2e6"), spaceAfter=8, spaceBefore=6))

    # Locadora
    if CNPJ_LOJA or ENDERECO_LOJA:
        story.append(Paragraph("LOCADORA", ParagraphStyle("sec4", parent=styles["Normal"],
            fontName="Helvetica-Bold", fontSize=8, textColor=accent, spaceBefore=0, spaceAfter=4)))
        info_loja = f"{NOME_LOJA}"
        if CNPJ_LOJA:   info_loja += f"  ·  CNPJ {CNPJ_LOJA}"
        if ENDERECO_LOJA: info_loja += f"  ·  {ENDERECO_LOJA}"
        story.append(Paragraph(info_loja, value_sm))
        story.append(Spacer(1, 12))

    # Nota de assinatura
    story.append(Paragraph(
        "Ao assinar digitalmente este documento, o LOCATÁRIO declara ter lido, "
        "compreendido e aceito todas as cláusulas e condições do contrato anexo.",
        ParagraphStyle("nota", parent=styles["Normal"],
            fontName="Helvetica-Oblique", fontSize=8.5, textColor=muted,
            spaceBefore=0, spaceAfter=0, alignment=TA_CENTER)))

    def rodape(canvas, doc):
        canvas.saveState()
        canvas.setFont("Helvetica-Oblique", 6.5)
        canvas.setFillColor(muted)
        canvas.drawString(2*cm, 1.2*cm,
            f"Contrato #{loc['id']:05d}  |  {NOME_LOJA}  |  "
            "Gerado por ClicksZap  |  www.clickszap.com.br")
        canvas.restoreState()

    doc.build(story, onFirstPage=rodape, onLaterPages=rodape)
    return buf.getvalue()


def _preencher_pdf_formulario(pdf_bytes: bytes, locs) -> bytes:
    """
    Preenche os campos de formulário (AcroForm) de um PDF com os dados da locação.

    Os campos do PDF devem ter nome igual aos campos do sistema, com ou sem chaves:
    JOGO_1, NOME_CLIENTE, VALOR_LOCACAO, etc. (também aceita {{JOGO_1}}).
    A comparação é case-insensitive.

    Se o PDF não tiver campos de formulário, retorna o PDF original sem alteração.
    """
    from pypdf import PdfReader, PdfWriter

    try:
        reader = PdfReader(io.BytesIO(pdf_bytes))
    except Exception as e:
        logger.warning("[pdf-form] PDF inválido: %s", e)
        return pdf_bytes

    # PDF tem campos de formulário?
    try:
        fields = reader.get_fields()
    except Exception:
        fields = None
    if not fields:
        # Sem campos — devolve o PDF como está
        return pdf_bytes

    campos = _build_campos(locs)
    # Mapa case-insensitive, com e sem chaves: {"jogo_1": valor, "{{jogo_1}}": valor}
    valores = {}
    for nome, valor in campos.items():
        valores[nome.lower()] = valor
        valores["{{" + nome.lower() + "}}"] = valor

    # Monta dicionário {nome_real_do_campo: valor} para os campos que casarem
    preenchimento = {}
    for nome_campo in fields.keys():
        chave = (nome_campo or "").strip().lower()
        if chave in valores:
            preenchimento[nome_campo] = valores[chave]

    if not preenchimento:
        logger.info("[pdf-form] PDF tem campos mas nenhum casou com os do sistema")
        return pdf_bytes

    # A partir daqui, QUALQUER falha no preenchimento devolve o PDF original
    # — o contrato é enviado mesmo assim, só sem os campos preenchidos.
    try:
        writer = PdfWriter()
        writer.append(reader)

        # Preenche e "achata" os campos: os valores viram conteúdo fixo da página,
        # aparecendo em qualquer visualizador (inclusive o serviço de assinatura).
        achatou = True
        for page in writer.pages:
            try:
                writer.update_page_form_field_values(
                    page, preenchimento, auto_regenerate=False, flatten=True
                )
            except TypeError:
                # Versão antiga do pypdf sem flatten/auto_regenerate
                achatou = False
                writer.update_page_form_field_values(page, preenchimento)

        # Fallback p/ versões antigas: força o visualizador a regenerar a aparência
        if not achatou:
            try:
                writer.set_need_appearances_writer(True)
            except Exception:
                pass

        out = io.BytesIO()
        writer.write(out)
        logger.info("[pdf-form] %d campo(s) preenchido(s) no PDF", len(preenchimento))
        return out.getvalue()
    except Exception as e:
        logger.error("[pdf-form] Falha ao preencher — enviando PDF original: %s", e)
        return pdf_bytes


def gerar_pdf_contrato(locacao_id: int) -> bytes:
    """
    Gera o PDF final do contrato, cobrindo todos os jogos alugados pelo
    mesmo cliente no mesmo dia (grupo automático).

    Modo A — DOCX do usuário (prioridade máxima):
      Substitui {{CAMPOS}} numerados e gera PDF via ReportLab.

    Modo B — PDF base do usuário:
      Página 1: Folha de Dados (capa automática). Páginas 2+: PDF do usuário.

    Modo C — Apenas texto do editor:
      PDF gerado 100% do modelo de texto.
    """
    from pypdf import PdfWriter, PdfReader

    locs = _buscar_grupo_locacoes(locacao_id)
    if not locs:
        raise ValueError(f"Locação #{locacao_id} não encontrada")
    loc = locs[0]  # locação principal

    template_bytes = carregar_template_pdf()

    # ── Modo A: DOCX do usuário (prioridade máxima) ───────────────────────────
    docx_bytes = carregar_template_docx()
    if docx_bytes:
        return _docx_para_pdf(docx_bytes, locs)

    if template_bytes:
        # ── Modo B: PDF do usuário — preenche campos de formulário se houver ──
        return _preencher_pdf_formulario(template_bytes, locs)

    else:
        # ── Modo C: PDF gerado 100% do modelo de texto ────────────────────────
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.units import cm
            from reportlab.lib.colors import HexColor
            from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, HRFlowable, Image as RLImage)
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
            import re
        except ImportError:
            raise RuntimeError("reportlab não instalado. Execute: pip install reportlab")

        modelo_texto = carregar_modelo()
        texto_final  = _substituir_campos(modelo_texto, locs)

        dark   = HexColor("#1a1a2e")
        accent = HexColor("#ED940E")
        muted  = HexColor("#6c757d")
        txt    = HexColor("#212529")

        styles = getSampleStyleSheet()
        normal = ParagraphStyle("normal", parent=styles["Normal"],
            fontName="Helvetica", fontSize=9.5, leading=14,
            textColor=txt, spaceBefore=0, spaceAfter=4, alignment=TA_JUSTIFY)
        titulo = ParagraphStyle("titulo", parent=styles["Normal"],
            fontName="Helvetica-Bold", fontSize=11, leading=15,
            textColor=dark, spaceBefore=8, spaceAfter=2)
        assinatura = ParagraphStyle("assinatura", parent=styles["Normal"],
            fontName="Courier", fontSize=8.5, leading=13,
            textColor=muted, spaceBefore=4, spaceAfter=0, alignment=TA_CENTER)

        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=A4,
            leftMargin=2*cm, rightMargin=2*cm,
            topMargin=2.5*cm, bottomMargin=2.5*cm)
        story = []

        _logo = os.path.join(os.path.dirname(os.path.abspath(__file__)), "jogoteka_colorido.png")
        if os.path.exists(_logo):
            _img = RLImage(_logo, width=6*cm, height=3*cm, kind="proportional")
            _img.hAlign = "CENTER"
            story.append(_img)
            story.append(Spacer(1, 6))
        else:
            story.append(Paragraph(NOME_LOJA, ParagraphStyle("hdr", parent=styles["Normal"],
                fontName="Helvetica-Bold", fontSize=20, textColor=accent,
                spaceBefore=0, spaceAfter=2, alignment=TA_CENTER)))
        story.append(HRFlowable(width="100%", thickness=1.5, color=accent, spaceAfter=10))

        from reportlab.platypus import Table as _SigTable, TableStyle as _SigStyle
        s_sig = ParagraphStyle("sig", parent=assinatura, alignment=TA_CENTER)

        _lines = texto_final.split("\n")
        _i = 0
        while _i < len(_lines):
            linha = _lines[_i].rstrip()
            _i += 1

            if not linha:
                story.append(Spacer(1, 6)); continue

            # Títulos centralizados: # Texto → H1 grande bold  |  ## Texto → H2 médio
            if linha.startswith("## "):
                story.append(Paragraph(linha[3:], ParagraphStyle("h2c", parent=normal,
                    fontName="Helvetica-Bold", fontSize=12, leading=16,
                    textColor=dark, spaceBefore=4, spaceAfter=6,
                    alignment=TA_CENTER))); continue
            if linha.startswith("# "):
                story.append(Paragraph(linha[2:], ParagraphStyle("h1c", parent=normal,
                    fontName="Helvetica-Bold", fontSize=16, leading=20,
                    textColor=dark, spaceBefore=6, spaceAfter=4,
                    alignment=TA_CENTER))); continue

            # Dupla linha de assinatura: dois grupos de ____ separados por espaços
            if re.search(r'_{10,}\s+_{10,}', linha):
                prox = _lines[_i].rstrip() if _i < len(_lines) else ""
                tem_nomes = "Locadora" in prox or "Locatário" in prox or "Locatario" in prox
                if tem_nomes:
                    _i += 1  # consome a linha dos nomes junto
                    partes = [p.strip() for p in re.split(r'\s{2,}', prox) if p.strip()]
                else:
                    partes = []
                while len(partes) < 2:
                    partes.append("")
                esq, dir_ = partes[0], partes[-1]

                sig_data = [
                    [Paragraph("______________________________", s_sig),
                     Paragraph("______________________________", s_sig)],
                    [Paragraph(esq,  s_sig),
                     Paragraph(dir_, s_sig)],
                ]
                t_sig = _SigTable(sig_data, colWidths=["50%", "50%"])
                t_sig.setStyle(_SigStyle([
                    ("VALIGN",        (0, 0), (-1, -1), "TOP"),
                    ("TOPPADDING",    (0, 0), (-1, -1), 2),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ]))
                story.append(Spacer(1, 20))
                story.append(t_sig)
                continue

            if linha.startswith("___"):
                story.append(HRFlowable(width="45%", thickness=0.5, color=muted,
                                        spaceAfter=2, spaceBefore=8)); continue
            if re.match(r"^\d+\.\s+[A-ZÁÉÍÓÚ]", linha):
                story.append(Spacer(1, 6))
                story.append(Paragraph(linha, titulo))
                story.append(HRFlowable(width="100%", thickness=0.3,
                                        color=HexColor("#dee2e6"), spaceAfter=4)); continue
            if re.match(r"^\d+\.\d+\s", linha):
                story.append(Paragraph(linha, ParagraphStyle(
                    "subitem", parent=normal, leftIndent=10, spaceBefore=4))); continue
            if "_____" in linha or ("Locadora" in linha and "Locatário" in linha):
                story.append(Spacer(1, 12))
                story.append(Paragraph(linha, assinatura)); continue
            story.append(Paragraph(linha, normal))

        def _rodape(canvas, doc):
            canvas.saveState()
            canvas.setFont("Helvetica-Oblique", 6.5)
            canvas.setFillColor(muted)
            canvas.drawString(2*cm, 1.2*cm,
                f"Contrato #{loc['id']:05d}  |  {NOME_LOJA}  |  "
                "Gerado por ClicksZap  |  www.clickszap.com.br")
            canvas.drawRightString(A4[0] - 2*cm, 1.2*cm, f"Página {doc.page}")
            canvas.restoreState()

        doc.build(story, onFirstPage=_rodape, onLaterPages=_rodape)
        return buf.getvalue()


# ── Integração ClicksZap ──────────────────────────────────────────────────────

def _httpx():
    """Import lazy do httpx para não quebrar o startup se não instalado."""
    try:
        import httpx as _hx
        return _hx
    except ImportError:
        raise RuntimeError("httpx não instalado. Execute: pip install httpx")


def _bearer_auth():
    """
    Auth customizado que reaplica o Bearer token em cada redirect.
    O httpx padrão descarta o Authorization em cross-origin redirects,
    causando 403 no ClicksZap.
    """
    hx = _httpx()

    class _BearerAuth(hx.Auth):
        def auth_flow(self, request):
            request.headers["Authorization"] = f"Bearer {_get_token()}"
            yield request

    return _BearerAuth()


def enviar_contrato(locacao_id: int) -> dict:
    """
    Gera o PDF, faz upload no ClicksZap e cria a solicitação de assinatura.
    Salva contrato_token, contrato_status e contrato_request_id na locação.
    """
    token_atual = _get_token()
    if not token_atual:
        return {"error": "CLICKSZAP_TOKEN não configurado. Acesse www.clickszap.com.br/panel/api-token e adicione no .env do Render."}

    locs = _buscar_grupo_locacoes(locacao_id)
    if not locs:
        return {"error": "Locação não encontrada"}
    loc = locs[0]

    cliente_nome = loc.get("cliente_nome") or "Cliente"
    cliente_tel  = re.sub(r'\D', '', loc.get("cliente_tel") or "")
    if not cliente_tel:
        return {"error": "Cliente sem telefone cadastrado — adicione o telefone antes de enviar o contrato"}
    if not cliente_tel.startswith("55"):
        cliente_tel = "55" + cliente_tel

    n_jogos = len(locs)
    nome_pdf = (f"Contrato_{loc['id']:05d}.pdf" if n_jogos == 1
                else f"Contrato_Grupo_{loc['id']:05d}_{n_jogos}jogos.pdf")

    # 1. Gera PDF com o modelo + dados substituídos (grupo completo)
    try:
        pdf_bytes = gerar_pdf_contrato(locacao_id)
    except Exception as e:
        logger.error("[contratos] Erro ao gerar PDF: %s", e)
        return {"error": f"Erro ao gerar contrato: {e}"}

    hx = _httpx()

    tok = _get_token()
    logger.info("[contratos] TOKEN carregado: %s... (len=%d)", tok[:12] if tok else "VAZIO", len(tok))
    logger.info("[contratos] URL base: %s", CLICKSZAP_URL)

    auth = _bearer_auth()

    # 2. Upload do PDF no ClicksZap
    try:
        upload_url = f"{CLICKSZAP_URL}/documents"
        logger.info("[contratos] POST %s (PDF %d bytes)", upload_url, len(pdf_bytes))
        resp = hx.post(
            upload_url,
            auth=auth,
            files={"file": (nome_pdf, pdf_bytes, "application/pdf")},
            timeout=30,
            follow_redirects=True,
        )
        logger.info("[contratos] Upload resposta: %s — %s", resp.status_code, resp.text[:300])
        if resp.status_code != 201:
            logger.error("[contratos] Upload falhou: %s %s", resp.status_code, resp.text)
            return {"error": f"Falha no upload do PDF no ClicksZap ({resp.status_code}): {resp.text[:200]}"}
        document_id = resp.json()["id"]
    except Exception as e:
        logger.error("[contratos] Erro no upload: %s", e)
        return {"error": f"Erro de conexão com ClicksZap: {e}"}

    # 3. Cria solicitação de assinatura (dispara WhatsApp automaticamente)
    try:
        resp = hx.post(
            f"{CLICKSZAP_URL}/signature-requests",
            auth=auth,
            json={
                "document_id": document_id,
                "signer_name": cliente_nome,
                "signer_phone": cliente_tel,
            },
            timeout=30,
            follow_redirects=True,
        )
        if resp.status_code != 201:
            logger.error("[contratos] Signature request falhou: %s %s", resp.status_code, resp.text)
            return {"error": f"Falha ao criar solicitação de assinatura ({resp.status_code}): {resp.text[:200]}"}
        data = resp.json()
    except Exception as e:
        logger.error("[contratos] Erro ao criar signature request: %s", e)
        return {"error": f"Erro de conexão com ClicksZap: {e}"}

    request_id   = str(data["id"])
    token        = data["token"]
    signing_link = f"{CLICKSZAP_URL}/s/{token}"

    # 4. Salva em TODAS as locações do grupo
    _salvar_dados_contrato([l["id"] for l in locs], request_id, token, "pending")

    return {
        "ok":          True,
        "request_id":  request_id,
        "token":       token,
        "signing_link": signing_link,
    }


# ── Lembretes de devolução ────────────────────────────────────────────────────

TEMPLATE_LEMBRETE_PADRAO = (
    "Olá *{nome}*! 👋\n\n"
    "Passando para lembrar que o jogo *{jogo}* deve ser devolvido *amanhã, dia {data}*. 🎲\n\n"
    "Qualquer dúvida é só chamar. Obrigado! 😊\n\n"
    "— Jogoteka 🎲"
)

def carregar_template_lembrete() -> str:
    try:
        with get_connection() as conn:
            row = conn.execute(
                "SELECT template FROM mensagem_lembrete WHERE ativo = 1 ORDER BY id DESC LIMIT 1"
            ).fetchone()
            if row and row["template"]:
                return row["template"]
    except Exception:
        pass
    return TEMPLATE_LEMBRETE_PADRAO

def salvar_template_lembrete(template: str, nome: str = "Lembrete de Devolução"):
    agora = datetime.now().isoformat(sep=" ", timespec="seconds")
    with get_connection() as conn:
        existing = conn.execute("SELECT id FROM mensagem_lembrete LIMIT 1").fetchone()
        if existing:
            conn.execute(
                "UPDATE mensagem_lembrete SET template = ?, nome = ?, atualizado_em = ? WHERE id = ?",
                (template, nome, agora, existing["id"])
            )
        else:
            conn.execute(
                "INSERT INTO mensagem_lembrete (nome, template, ativo, atualizado_em) VALUES (?, ?, 1, ?)",
                (nome, template, agora)
            )

def _enviar_mensagem_whatsapp(telefone: str, mensagem: str) -> dict:
    """Envia mensagem de texto simples via ClicksZap."""
    token_atual = _get_token()
    if not token_atual:
        return {"error": "CLICKSZAP_TOKEN não configurado"}
    hx = _httpx()
    try:
        resp = hx.post(
            f"{CLICKSZAP_URL}/messages",
            auth=_bearer_auth(),
            json={"to": telefone, "text": mensagem},
            timeout=15,
            follow_redirects=True,
        )
        logger.info("[lembrete] POST /messages → %s: %s", resp.status_code, resp.text[:200])
        if resp.status_code not in (200, 201):
            return {"error": f"ClicksZap retornou {resp.status_code}: {resp.text[:200]}"}
        return {"ok": True}
    except Exception as e:
        logger.error("[lembrete] Erro ao enviar mensagem: %s", e)
        return {"error": str(e)}

def enviar_lembretes_devolucao() -> dict:
    """
    Busca locações ativas com data_prevista = amanhã e envia lembrete via WhatsApp.
    Evita duplicatas verificando o lembretes_log do dia.
    """
    from datetime import date, timedelta
    amanha = (date.today() + timedelta(days=1)).isoformat()
    hoje   = date.today().isoformat()

    template = carregar_template_lembrete()

    with get_connection() as conn:
        locacoes = conn.execute("""
            SELECT l.id, l.data_prevista,
                   j.nome AS jogo_nome,
                   c.nome AS cliente_nome, c.telefone AS cliente_tel
            FROM locacoes l
            JOIN jogos j ON j.id = l.jogo_id
            LEFT JOIN clientes c ON c.id = l.cliente_id
            WHERE l.status = 'ativa'
              AND substr(l.data_prevista, 1, 10) = ?
              AND c.telefone IS NOT NULL AND c.telefone != ''
        """, (amanha,)).fetchall()

        # Locações já notificadas hoje (evita duplicata em caso de restart)
        ja_enviados = set(
            row[0] for row in conn.execute(
                "SELECT locacao_id FROM lembretes_log WHERE status='ok' AND substr(enviado_em,1,10)=?",
                (hoje,)
            ).fetchall()
        )

    enviados = erros = 0
    agora_str = datetime.now().isoformat(sep=" ", timespec="seconds")

    try:
        data_fmt = datetime.strptime(amanha, "%Y-%m-%d").strftime("%d/%m")
    except Exception:
        data_fmt = amanha

    for loc in locacoes:
        if loc["id"] in ja_enviados:
            logger.info("[lembrete] Locação #%d já notificada hoje, pulando.", loc["id"])
            continue

        mensagem = (template
            .replace("{nome}", loc["cliente_nome"] or "Cliente")
            .replace("{jogo}", loc["jogo_nome"] or "jogo")
            .replace("{data}", data_fmt)
        )

        tel = (loc["cliente_tel"] or "").replace(" ","").replace("-","").replace("(","").replace(")","")
        if not tel.startswith("55"):
            tel = "55" + tel

        resultado = _enviar_mensagem_whatsapp(tel, mensagem)

        with get_connection() as conn:
            conn.execute(
                """INSERT INTO lembretes_log
                   (locacao_id, cliente_nome, cliente_tel, jogo_nome, data_prevista, status, erro, enviado_em)
                   VALUES (?,?,?,?,?,?,?,?)""",
                (loc["id"], loc["cliente_nome"], loc["cliente_tel"],
                 loc["jogo_nome"], loc["data_prevista"],
                 "ok" if resultado.get("ok") else "erro",
                 resultado.get("error"), agora_str)
            )

        if resultado.get("ok"):
            enviados += 1
            logger.info("[lembrete] ✔ Enviado para %s (locação #%d)", loc["cliente_nome"], loc["id"])
        else:
            erros += 1
            logger.error("[lembrete] ✘ Erro locação #%d: %s", loc["id"], resultado.get("error"))

    logger.info("[lembrete] Rodada concluída — enviados: %d | erros: %d | total: %d",
                enviados, erros, len(locacoes))
    return {"enviados": enviados, "erros": erros, "total": len(locacoes)}


def status_contrato(locacao_id: int) -> dict:
    """Consulta o ClicksZap pelo status do contrato da locação."""
    if not _get_token():
        return {"error": "CLICKSZAP_TOKEN não configurado"}

    with get_connection() as conn:
        row = conn.execute(
            "SELECT contrato_token, contrato_status, contrato_request_id FROM locacoes WHERE id = ?",
            (locacao_id,)
        ).fetchone()

    if not row:
        return {"error": "Locação não encontrada"}

    token      = row["contrato_token"]
    request_id = row["contrato_request_id"]

    if not token:
        return {"status": None, "message": "Contrato ainda não enviado"}

    hx = _httpx()
    try:
        resp = hx.get(
            f"{CLICKSZAP_URL}/signature-requests/{request_id}",
            auth=_bearer_auth(),
            timeout=15,
            follow_redirects=True,
        )
        if resp.status_code != 200:
            return {"error": f"Falha ao consultar ClicksZap ({resp.status_code})"}
        data = resp.json()
    except Exception as e:
        return {"error": f"Erro de conexão: {e}"}

    status = data.get("status", "pending")
    if status != row["contrato_status"]:
        _atualizar_status_contrato(locacao_id, status)

    return {
        "status":       status,
        "token":        token,
        "signing_link": f"{CLICKSZAP_URL}/s/{token}",
        "download_url": f"{CLICKSZAP_URL}/s/{token}/download" if status == "signed" else None,
        "signed_at":    data.get("signed_at"),
    }
